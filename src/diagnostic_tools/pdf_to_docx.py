#!/usr/bin/env python3
"""
PDF to DOCX Converter - Multiple Approaches
This script demonstrates different methods to convert PDF files to DOCX format.
"""

import os
import sys
from pathlib import Path


# Method 1: Using pdf2docx (Best for preserving layout and formatting)
def convert_with_pdf2docx(pdf_path, docx_path):
    """
    Convert PDF to DOCX using pdf2docx library.
    This method preserves formatting, tables, and images better.

    Install: pip install pdf2docx
    """
    try:
        from pdf2docx import Converter

        print(f"Converting {pdf_path} to {docx_path} using pdf2docx...")

        # Create converter object
        cv = Converter(pdf_path)

        # Convert with options
        cv.convert(docx_path, start=0, end=None)  # Convert all pages
        cv.close()

        print(f"✓ Conversion completed: {docx_path}")
        return True

    except ImportError:
        print("❌ pdf2docx not installed. Run: pip install pdf2docx")
        return False
    except Exception as e:
        print(f"❌ Error with pdf2docx: {e}")
        return False


# Method 2: Using PyPDF2 + python-docx (Text only, no formatting)
def convert_with_pypdf2(pdf_path, docx_path):
    """
    Convert PDF to DOCX using PyPDF2 for text extraction and python-docx for creation.
    This method only preserves text, no formatting or images.

    Install: pip install PyPDF2 python-docx
    """
    try:
        import PyPDF2
        from docx import Document

        print(f"Converting {pdf_path} to {docx_path} using PyPDF2...")

        # Extract text from PDF
        text_content = []
        with open(pdf_path, "rb") as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text.strip():
                    text_content.append(f"--- Page {page_num + 1} ---\n{text}\n")

        # Create DOCX document
        doc = Document()
        doc.add_heading("PDF Conversion", level=1)

        for content in text_content:
            doc.add_paragraph(content)

        doc.save(docx_path)
        print(f"✓ Text-only conversion completed: {docx_path}")
        return True

    except ImportError:
        print(
            "❌ Required libraries not installed. Run: pip install PyPDF2 python-docx"
        )
        return False
    except Exception as e:
        print(f"❌ Error with PyPDF2 method: {e}")
        return False


# Method 3: Using pdfplumber + python-docx (Better text extraction)
def convert_with_pdfplumber(pdf_path, docx_path):
    """
    Convert PDF to DOCX using pdfplumber for text extraction.
    Better at handling tables and complex layouts than PyPDF2.

    Install: pip install pdfplumber python-docx
    """
    try:
        import pdfplumber
        from docx import Document
        from docx.shared import Inches

        print(f"Converting {pdf_path} to {docx_path} using pdfplumber...")

        doc = Document()
        doc.add_heading("PDF Conversion with pdfplumber", level=1)

        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # Add page header
                doc.add_heading(f"Page {page_num + 1}", level=2)

                # Extract text
                text = page.extract_text()
                if text:
                    doc.add_paragraph(text)

                # Extract tables if any
                tables = page.extract_tables()
                if tables:
                    doc.add_paragraph("Tables found on this page:")
                    for table in tables:
                        # Add table to document
                        docx_table = doc.add_table(
                            rows=len(table), cols=len(table[0]) if table else 0
                        )
                        for i, row in enumerate(table):
                            for j, cell in enumerate(row):
                                if cell:
                                    docx_table.cell(i, j).text = str(cell)

                doc.add_page_break()

        doc.save(docx_path)
        print(f"✓ Enhanced conversion completed: {docx_path}")
        return True

    except ImportError:
        print(
            "❌ Required libraries not installed. Run: pip install pdfplumber python-docx"
        )
        return False
    except Exception as e:
        print(f"❌ Error with pdfplumber method: {e}")
        return False


def main():
    """Main function to demonstrate PDF to DOCX conversion."""

    # Check command line arguments
    if len(sys.argv) != 2:
        print("Usage: python pdf_to_docx.py <pdf_file_path>")
        print("Example: python pdf_to_docx.py document.pdf")
        return

    pdf_path = sys.argv[1]

    # Verify PDF file exists
    if not os.path.exists(pdf_path):
        print(f"❌ Error: PDF file '{pdf_path}' not found.")
        return

    # Create output filename
    pdf_name = Path(pdf_path).stem
    output_dir = Path(pdf_path).parent

    # Try different conversion methods
    methods = [
        ("pdf2docx", convert_with_pdf2docx, f"{output_dir}/{pdf_name}_pdf2docx.docx"),
        (
            "pdfplumber",
            convert_with_pdfplumber,
            f"{output_dir}/{pdf_name}_pdfplumber.docx",
        ),
        ("PyPDF2", convert_with_pypdf2, f"{output_dir}/{pdf_name}_pypdf2.docx"),
    ]

    print(f"🔄 Converting PDF: {pdf_path}")
    print("=" * 60)

    successful_conversions = 0

    for method_name, method_func, output_path in methods:
        print(f"\n📝 Trying method: {method_name}")
        if method_func(pdf_path, output_path):
            successful_conversions += 1
        print("-" * 40)

    print(f"\n📊 Summary: {successful_conversions}/{len(methods)} methods succeeded")

    if successful_conversions > 0:
        print("\n💡 Comparison of methods:")
        print("• pdf2docx: Best formatting preservation, handles images/tables")
        print("• pdfplumber: Good for tables, moderate formatting")
        print("• PyPDF2: Basic text only, fastest but limited")
    else:
        print("\n❌ No conversions succeeded. Install required libraries:")
        print("pip install pdf2docx pdfplumber PyPDF2 python-docx")


if __name__ == "__main__":
    main()


# Additional utility functions
def batch_convert_directory(directory_path, method="pdf2docx"):
    """Convert all PDF files in a directory."""
    pdf_files = list(Path(directory_path).glob("*.pdf"))

    if not pdf_files:
        print(f"No PDF files found in {directory_path}")
        return

    print(f"Found {len(pdf_files)} PDF files to convert...")

    for pdf_file in pdf_files:
        output_file = pdf_file.with_suffix(".docx")

        if method == "pdf2docx":
            convert_with_pdf2docx(str(pdf_file), str(output_file))
        elif method == "pdfplumber":
            convert_with_pdfplumber(str(pdf_file), str(output_file))
        elif method == "pypdf2":
            convert_with_pypdf2(str(pdf_file), str(output_file))


def install_requirements():
    """Install required packages."""
    import subprocess

    packages = ["pdf2docx", "pdfplumber", "PyPDF2", "python-docx"]

    for package in packages:
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", package])
            print(f"✓ Installed {package}")
        except subprocess.CalledProcessError:
            print(f"❌ Failed to install {package}")


# Example usage:
"""
# Basic usage:
python pdf_to_docx.py document.pdf

# Batch convert all PDFs in a directory:
batch_convert_directory("/path/to/pdfs/", method="pdf2docx")

# Install requirements:
install_requirements()
"""
