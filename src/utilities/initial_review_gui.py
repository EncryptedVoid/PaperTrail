"""
Multi-File Type Preview and Sorter GUI

Required packages:
pip install pillow PyMuPDF opencv-python python-docx ebooklib beautifulsoup4

Usage:
    Edit the variables below, then run: python file_sorter.py
"""

import os
import shutil
import threading
import tkinter as tk
from pathlib import Path
from tkinter import filedialog, messagebox, scrolledtext

import cv2
import fitz  # PyMuPDF
from PIL import Image, ImageTk

# Optional imports for better document support
try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from ebooklib import epub
    from bs4 import BeautifulSoup

    EPUB_LIB_AVAILABLE = True
except ImportError:
    EPUB_LIB_AVAILABLE = False

# Suppress MuPDF warnings
fitz.TOOLS.mupdf_display_errors(False)


class FilePreviewApp:
    def __init__(self, root, folder_path=None, destinations=None, labels=None):
        self.root = root
        self.root.title("File Preview and Sorter")
        self.root.geometry("1200x800")

        self.current_index = 0
        self.files = []
        self.video_thread = None
        self.video_playing = False
        self.video_cap = None

        # Sorting destinations and labels
        self.destinations = destinations or [None, None, None, None]
        self.labels = labels or ["Category 1", "Category 2", "Category 3", "Category 4"]

        # Setup UI
        self.setup_ui()

        # Load files if folder provided
        if folder_path and os.path.isdir(folder_path):
            self.load_folder(folder_path)

        # Bind keyboard events
        self.root.bind("<Left>", lambda e: self.previous_file())
        self.root.bind("<Right>", lambda e: self.next_file())
        self.root.bind("1", lambda e: self.move_file_to_category(0))
        self.root.bind("2", lambda e: self.move_file_to_category(1))
        self.root.bind("3", lambda e: self.move_file_to_category(2))
        self.root.bind("4", lambda e: self.move_file_to_category(3))
        self.root.bind("<Escape>", lambda e: self.stop_video())

    def setup_ui(self):
        # Top frame for controls
        top_frame = tk.Frame(self.root)
        top_frame.pack(side=tk.TOP, fill=tk.X, padx=5, pady=5)

        tk.Button(top_frame, text="Select Folder", command=self.select_folder).pack(
            side=tk.LEFT, padx=5
        )

        self.info_label = tk.Label(top_frame, text="No folder loaded")
        self.info_label.pack(side=tk.LEFT, padx=20)

        # File type display (large, bold, on right)
        self.file_type_label = tk.Label(
            top_frame,
            text="",
            font=("Arial", 20, "bold"),
            bg="lightgray",
            padx=15,
            pady=5,
            relief=tk.RAISED,
            borderwidth=2,
        )
        self.file_type_label.pack(side=tk.RIGHT, padx=10)

        self.file_counter = tk.Label(top_frame, text="")
        self.file_counter.pack(side=tk.RIGHT, padx=5)

        # Main preview area
        self.preview_frame = tk.Frame(self.root, bg="gray20")
        self.preview_frame.pack(side=tk.TOP, fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Bottom frame with action buttons and instructions
        bottom_frame = tk.Frame(self.root, bg="lightblue")
        bottom_frame.pack(side=tk.BOTTOM, fill=tk.X)

        # Navigation instructions
        nav_frame = tk.Frame(bottom_frame, bg="lightblue")
        nav_frame.pack(side=tk.TOP, fill=tk.X, pady=5)
        tk.Label(
            nav_frame,
            text="Navigation: ← Previous | → Next",
            bg="lightblue",
            font=("Arial", 10, "bold"),
        ).pack()

        # Action buttons frame
        buttons_frame = tk.Frame(bottom_frame, bg="lightblue")
        buttons_frame.pack(side=tk.TOP, fill=tk.X, padx=20, pady=5)

        # Create 4 action buttons
        button_colors = ["#90EE90", "#FFB6C1", "#87CEEB", "#FFD700"]
        for i in range(4):
            btn_frame = tk.Frame(buttons_frame, bg="lightblue")
            btn_frame.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=5)

            btn = tk.Button(
                btn_frame,
                text=f"[{i+1}] {self.labels[i]}",
                command=lambda idx=i: self.move_file_to_category(idx),
                font=("Arial", 12, "bold"),
                bg=button_colors[i],
                padx=20,
                pady=10,
                relief=tk.RAISED,
                borderwidth=3,
            )
            btn.pack(fill=tk.X)

            # Show destination path if available
            if self.destinations[i]:
                dest_label = tk.Label(
                    btn_frame,
                    text=f"→ {self.destinations[i]}",
                    bg="lightblue",
                    font=("Arial", 8),
                    fg="gray30",
                )
                dest_label.pack()

    def select_folder(self):
        folder = filedialog.askdirectory()
        if folder:
            self.load_folder(folder)

    def load_folder(self, folder_path):
        self.files = []
        for root, dirs, files in os.walk(folder_path):
            for file in files:
                self.files.append(os.path.join(root, file))

        self.files.sort()
        self.current_index = 0

        if self.files:
            self.info_label.config(text=f"Loaded {len(self.files)} files")
            self.show_current_file()
        else:
            self.info_label.config(text="No files found in folder")

    def show_current_file(self):
        if not self.files or self.current_index >= len(self.files):
            return

        # Stop any playing video
        self.stop_video()

        # Clear preview frame
        for widget in self.preview_frame.winfo_children():
            widget.destroy()

        file_path = self.files[self.current_index]
        file_name = os.path.basename(file_path)
        ext = Path(file_path).suffix.lower()

        # Update counter
        self.file_counter.config(
            text=f"File {self.current_index + 1}/{len(self.files)}: {file_name}"
        )

        # Update file type display
        if ext:
            self.file_type_label.config(text=ext.upper())
        else:
            self.file_type_label.config(text="NO EXT")

        # Determine file type and display
        try:
            if ext in [".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp"]:
                self.show_image(file_path)
            elif ext in [".pdf"]:
                self.show_pdf_document(file_path)
            elif ext in [".epub"]:
                self.show_epub_document(file_path)
            elif ext in [".docx", ".doc"]:
                self.show_docx_document(file_path)
            elif ext in [".pub"]:
                self.show_document_fallback(file_path)
            elif ext in [".mp4", ".avi", ".mov", ".mkv", ".webm", ".flv"]:
                self.show_video(file_path)
            elif ext in [".html", ".htm"]:
                self.show_html(file_path)
            elif ext in [
                ".txt",
                ".py",
                ".js",
                ".java",
                ".cpp",
                ".c",
                ".h",
                ".css",
                ".json",
                ".xml",
                ".md",
                ".csv",
                ".log",
                ".sh",
                ".yml",
                ".yaml",
            ]:
                self.show_text(file_path)
            else:
                self.show_text(file_path)  # Try to show as text
        except Exception as e:
            self.show_error(f"Error loading file: {str(e)}")

    def show_image(self, file_path):
        img = Image.open(file_path)

        # Resize to fit preview area
        preview_width = self.preview_frame.winfo_width() - 20
        preview_height = self.preview_frame.winfo_height() - 20

        if preview_width > 10 and preview_height > 10:
            img.thumbnail((preview_width, preview_height), Image.Resampling.LANCZOS)

        photo = ImageTk.PhotoImage(img)

        label = tk.Label(self.preview_frame, image=photo, bg="gray20")
        label.image = photo  # Keep a reference
        label.pack(expand=True)

    def show_pdf_document(self, file_path):
        """Display PDF documents (first 5 pages only)"""
        try:
            pdf_document = fitz.open(file_path)
        except Exception as e:
            self.show_error(f"Cannot open PDF: {str(e)}")
            return

        total_pages = len(pdf_document)
        pages_to_show = min(5, total_pages)

        # Create scrollable canvas
        canvas = tk.Canvas(self.preview_frame, bg="white")
        scrollbar = tk.Scrollbar(
            self.preview_frame, orient="vertical", command=canvas.yview
        )
        scrollable_frame = tk.Frame(canvas, bg="white")

        scrollable_frame.bind(
            "<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        # Show notice if PDF has more pages
        if total_pages > 5:
            notice = tk.Label(
                scrollable_frame,
                text=f"📄 Showing first 5 pages of {total_pages} total pages",
                bg="lightyellow",
                font=("Arial", 11, "bold"),
                pady=10,
            )
            notice.pack(fill=tk.X, pady=5)

        # Render first 5 pages only
        for page_num in range(pages_to_show):
            page = pdf_document[page_num]
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            photo = ImageTk.PhotoImage(img)

            page_label = tk.Label(scrollable_frame, image=photo, bg="white")
            page_label.image = photo  # Keep reference
            page_label.pack(pady=5)

            # Page number
            tk.Label(
                scrollable_frame,
                text=f"Page {page_num + 1}",
                bg="white",
                font=("Arial", 10, "bold"),
            ).pack()

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        pdf_document.close()

    def show_epub_document(self, file_path):
        """Display EPUB documents"""
        if not EPUB_LIB_AVAILABLE:
            # Fallback to PyMuPDF for EPUB
            try:
                doc = fitz.open(file_path)
                self._render_fitz_document(doc)
                doc.close()
            except Exception as e:
                self.show_error(
                    f"Cannot open EPUB (install ebooklib for better support):\n{str(e)}"
                )
            return

        try:
            book = epub.read_epub(file_path)

            # Create scrollable text area
            text_widget = scrolledtext.ScrolledText(
                self.preview_frame, wrap=tk.WORD, font=("Georgia", 11), padx=20, pady=10
            )
            text_widget.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

            # Extract and display text from all chapters
            for item in book.get_items():
                if item.get_type() == 9:  # ITEM_DOCUMENT type
                    content = item.get_content().decode("utf-8", errors="ignore")

                    # Parse HTML and extract text
                    try:
                        soup = BeautifulSoup(content, "html.parser")
                        text = soup.get_text()
                        text_widget.insert(tk.END, text + "\n\n")
                    except:
                        text_widget.insert(tk.END, content + "\n\n")

            text_widget.config(state=tk.DISABLED)

        except Exception as e:
            self.show_error(f"Cannot open EPUB: {str(e)}")

    def show_docx_document(self, file_path):
        """Display DOCX documents"""
        if not DOCX_AVAILABLE:
            self.show_error(
                f"Cannot open DOCX files.\nInstall python-docx: pip install python-docx"
            )
            return

        try:
            doc = Document(file_path)

            # Create scrollable text area
            text_widget = scrolledtext.ScrolledText(
                self.preview_frame, wrap=tk.WORD, font=("Arial", 11), padx=20, pady=10
            )
            text_widget.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

            # Extract all paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_widget.insert(tk.END, para.text + "\n\n")

            # Extract tables if any
            if doc.tables:
                text_widget.insert(tk.END, "\n--- TABLES ---\n\n")
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text for cell in row.cells])
                        text_widget.insert(tk.END, row_text + "\n")
                    text_widget.insert(tk.END, "\n")

            text_widget.config(state=tk.DISABLED)

        except Exception as e:
            self.show_error(f"Cannot open DOCX: {str(e)}")

    def show_document_fallback(self, file_path):
        """Fallback for documents that can't be opened with specialized libraries"""
        try:
            doc = fitz.open(file_path)
            self._render_fitz_document(doc)
            doc.close()
        except Exception as e:
            self.show_error(f"Cannot open document.\nTrying text view...")
            try:
                self.show_text(file_path)
            except:
                self.show_error(f"Cannot display this file type")

    def _render_fitz_document(self, doc):
        """Helper method to render any fitz document"""
        # Create scrollable canvas
        canvas = tk.Canvas(self.preview_frame, bg="white")
        scrollbar = tk.Scrollbar(
            self.preview_frame, orient="vertical", command=canvas.yview
        )
        scrollable_frame = tk.Frame(canvas, bg="white")

        scrollable_frame.bind(
            "<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        # Render all pages
        for page_num in range(len(doc)):
            page = doc[page_num]
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            photo = ImageTk.PhotoImage(img)

            page_label = tk.Label(scrollable_frame, image=photo, bg="white")
            page_label.image = photo  # Keep reference
            page_label.pack(pady=5)

            # Page number
            tk.Label(
                scrollable_frame,
                text=f"Page {page_num + 1}",
                bg="white",
                font=("Arial", 10, "bold"),
            ).pack()

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

    def show_video(self, file_path):
        # Create video display label
        video_label = tk.Label(self.preview_frame, bg="gray20")
        video_label.pack(expand=True)

        # Start video playback in separate thread
        self.video_playing = True
        self.video_cap = cv2.VideoCapture(file_path)

        # Mute by not playing audio (OpenCV doesn't handle audio)

        def play_video():
            while self.video_playing and self.video_cap.isOpened():
                ret, frame = self.video_cap.read()
                if not ret:
                    # Loop video
                    self.video_cap.set(cv2.CAP_PROP_POS_FRAMES, 0)
                    continue

                # Convert BGR to RGB
                frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                img = Image.fromarray(frame)

                # Resize to fit
                preview_width = self.preview_frame.winfo_width() - 20
                preview_height = self.preview_frame.winfo_height() - 20
                if preview_width > 10 and preview_height > 10:
                    img.thumbnail(
                        (preview_width, preview_height), Image.Resampling.LANCZOS
                    )

                photo = ImageTk.PhotoImage(img)

                if self.video_playing:
                    video_label.config(image=photo)
                    video_label.image = photo

                # Control frame rate (30 fps)
                self.root.after(33)

        self.video_thread = threading.Thread(target=play_video, daemon=True)
        self.video_thread.start()

    def show_html(self, file_path):
        # For simplicity, render HTML as text with a webview note
        # To properly render, you'd need tkinterweb or similar
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        text_widget = scrolledtext.ScrolledText(
            self.preview_frame, wrap=tk.WORD, font=("Courier", 10)
        )
        text_widget.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        text_widget.insert(1.0, f"HTML SOURCE:\n\n{content}")
        text_widget.config(state=tk.DISABLED)

    def show_text(self, file_path):
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
        except Exception as e:
            content = f"Could not read file as text: {str(e)}"

        text_widget = scrolledtext.ScrolledText(
            self.preview_frame, wrap=tk.WORD, font=("Courier", 10)
        )
        text_widget.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        text_widget.insert(1.0, content)
        text_widget.config(state=tk.DISABLED)

    def show_error(self, message):
        label = tk.Label(
            self.preview_frame, text=message, bg="gray20", fg="red", font=("Arial", 14)
        )
        label.pack(expand=True)

    def stop_video(self):
        self.video_playing = False
        if self.video_cap:
            self.video_cap.release()
            self.video_cap = None

    def move_file_to_category(self, category_index):
        """Move current file to specified category directory"""
        if not self.files or self.current_index >= len(self.files):
            return

        destination = self.destinations[category_index]

        if not destination:
            messagebox.showerror(
                "Error", f"No destination set for {self.labels[category_index]}"
            )
            return

        # Create destination directory if it doesn't exist
        os.makedirs(destination, exist_ok=True)

        source_file = self.files[self.current_index]
        file_name = os.path.basename(source_file)
        dest_file = os.path.join(destination, file_name)

        # Handle duplicate filenames
        counter = 1
        base_name, ext = os.path.splitext(file_name)
        while os.path.exists(dest_file):
            new_name = f"{base_name}_{counter}{ext}"
            dest_file = os.path.join(destination, new_name)
            counter += 1

        try:
            shutil.move(source_file, dest_file)
            print(f"Moved: {source_file} -> {dest_file}")

            # Remove from list
            self.files.pop(self.current_index)

            # Update info
            self.info_label.config(text=f"{len(self.files)} files remaining")

            # Show next file (or previous if at end)
            if self.files:
                if self.current_index >= len(self.files):
                    self.current_index = len(self.files) - 1
                self.show_current_file()
            else:
                self.info_label.config(text="All files sorted!")
                for widget in self.preview_frame.winfo_children():
                    widget.destroy()
                label = tk.Label(
                    self.preview_frame,
                    text="✓ All files sorted!",
                    bg="gray20",
                    fg="lightgreen",
                    font=("Arial", 24, "bold"),
                )
                label.pack(expand=True)

        except Exception as e:
            messagebox.showerror("Error", f"Failed to move file: {str(e)}")

    def next_file(self):
        """Navigate to next file"""
        if self.current_index < len(self.files) - 1:
            self.current_index += 1
            self.show_current_file()

    def previous_file(self):
        """Navigate to previous file"""
        if self.current_index > 0:
            self.current_index -= 1
            self.show_current_file()
