from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Union, Optional
import mimetypes
from contextlib import contextmanager


class MetadataExtractor:
    """Comprehensive file metadata extraction with lazy loading of dependencies"""

    def __init__(self):
        self._cache = {}

    def extract(self, filepath: Union[str, Path]) -> Dict[str, Any]:
        """Main entry point for metadata extraction"""
        filepath = Path(filepath)

        if not filepath.exists():
            return {"error": "File not found", "path": str(filepath)}

        # Always get basic info
        result = self._get_basic_info(filepath)

        # Route to specific extractors
        ext = filepath.suffix.lower()
        if ext in {
            ".jpg",
            ".jpeg",
            ".png",
            ".gif",
            ".bmp",
            ".tiff",
            ".tif",
            ".webp",
            ".heic",
            ".heif",
            ".raw",
            ".cr2",
            ".nef",
            ".arw",
        }:
            result.update(self._extract_image_metadata(filepath))
        elif ext in {".pdf", ".docx", ".xlsx", ".pptx"}:
            result.update(self._extract_document_metadata(filepath))
        else:
            result["warning"] = f"No specialized extractor for {ext}"

        return result

    def _get_basic_info(self, filepath: Path) -> Dict[str, Any]:
        """Extract filesystem metadata"""
        try:
            stat = filepath.stat()
            mime_type, _ = mimetypes.guess_type(str(filepath))

            return {
                "filename": filepath.name,
                "path": str(filepath.absolute()),
                "size_bytes": stat.st_size,
                "size_mb": round(stat.st_size / (1024 * 1024), 2),
                "created": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "extension": filepath.suffix.lower(),
                "mime_type": mime_type,
            }
        except Exception as e:
            return {"error": f"Failed to read basic info: {e}"}

    def _extract_image_metadata(self, filepath: Path) -> Dict[str, Any]:
        """Extract image metadata with lazy imports"""
        result = {"type": "image"}

        # PIL metadata
        pil_data = self._get_pil_metadata(filepath)
        if "error" not in pil_data:
            result.update(pil_data)

        # Advanced metadata (optional)
        advanced_data = self._get_advanced_image_metadata(filepath)
        if "error" not in advanced_data:
            result.update(advanced_data)

        return result

    def _get_pil_metadata(self, filepath: Path) -> Dict[str, Any]:
        """Extract PIL-based image metadata"""
        try:
            from PIL import Image
            from PIL.ExifTags import TAGS
            import pillow_heif

            pillow_heif.register_heif_opener()

            with Image.open(filepath) as img:
                # Basic info
                image_info = {
                    "dimensions": img.size,
                    "width": img.size[0],
                    "height": img.size[1],
                    "mode": img.mode,
                    "format": img.format,
                    "has_transparency": img.mode in ("RGBA", "LA")
                    or "transparency" in img.info,
                }

                # EXIF data
                exif_data = {}
                exif = img.getexif()  # Modern method, not deprecated _getexif
                if exif:
                    exif_data = {TAGS.get(k, f"tag_{k}"): v for k, v in exif.items()}

                return {"image_info": image_info, "exif": exif_data}

        except ImportError as e:
            return {"error": f"PIL not available: {e}"}
        except Exception as e:
            return {"error": f"PIL extraction failed: {e}"}

    def _get_advanced_image_metadata(self, filepath: Path) -> Dict[str, Any]:
        """Extract advanced metadata using pyexiv2 if available"""
        try:
            import pyexiv2

            with pyexiv2.Image(str(filepath)) as img:
                return {
                    "advanced_metadata": {
                        "exif": img.read_exif(),
                        "iptc": img.read_iptc(),
                        "xmp": img.read_xmp(),
                    }
                }

        except ImportError:
            return {"info": "pyexiv2 not available - using basic EXIF only"}
        except Exception as e:
            return {"error": f"Advanced metadata extraction failed: {e}"}

    def _extract_document_metadata(self, filepath: Path) -> Dict[str, Any]:
        """Route document extraction based on file type"""
        result = {"type": "document"}
        ext = filepath.suffix.lower()

        extractors = {
            ".pdf": self._extract_pdf_metadata,
            ".docx": self._extract_docx_metadata,
            ".xlsx": self._extract_xlsx_metadata,
            ".pptx": self._extract_pptx_metadata,
        }

        if ext in extractors:
            result.update(extractors[ext](filepath))
        else:
            result["error"] = f"No extractor for {ext}"

        return result

    def _extract_pdf_metadata(self, filepath: Path) -> Dict[str, Any]:
        """Extract PDF metadata - prefer PyMuPDF for comprehensive data"""
        # Try PyMuPDF first (best overall)
        pymupdf_result = self._try_pymupdf(filepath)
        if "error" not in pymupdf_result:
            return pymupdf_result

        # Fallback to pdfplumber for basic info
        return self._try_pdfplumber(filepath)

    def _try_pymupdf(self, filepath: Path) -> Dict[str, Any]:
        """PyMuPDF extraction with proper resource management"""
        try:
            import fitz

            doc = fitz.open(str(filepath))
            try:
                result = {
                    "pdf_info": {
                        "page_count": doc.page_count,
                        "is_encrypted": doc.needs_pass,
                        "metadata": doc.metadata,
                        "toc": doc.get_toc(),
                        "has_text": (
                            bool(doc[0].get_text()) if doc.page_count > 0 else False
                        ),
                    }
                }
                return result
            finally:
                doc.close()  # Ensure cleanup

        except ImportError:
            return {"error": "PyMuPDF not available"}
        except Exception as e:
            return {"error": f"PyMuPDF extraction failed: {e}"}

    def _try_pdfplumber(self, filepath: Path) -> Dict[str, Any]:
        """Fallback PDF extraction"""
        try:
            import pdfplumber

            with pdfplumber.open(filepath) as pdf:
                return {
                    "pdf_info": {
                        "page_count": len(pdf.pages),
                        "has_text": (
                            bool(pdf.pages[0].extract_text()) if pdf.pages else False
                        ),
                    }
                }
        except ImportError:
            return {"error": "No PDF libraries available"}
        except Exception as e:
            return {"error": f"PDF extraction failed: {e}"}

    def _extract_docx_metadata(self, filepath: Path) -> Dict[str, Any]:
        """Extract Word document metadata"""
        try:
            from docx import Document

            doc = Document(str(filepath))
            props = doc.core_properties

            return {
                "document_info": {
                    "title": props.title,
                    "author": props.author,
                    "created": props.created.isoformat() if props.created else None,
                    "modified": props.modified.isoformat() if props.modified else None,
                    "last_modified_by": props.last_modified_by,
                    "paragraph_count": len(doc.paragraphs),
                    "table_count": len(doc.tables),
                }
            }
        except ImportError:
            return {"error": "python-docx not available"}
        except Exception as e:
            return {"error": f"DOCX extraction failed: {e}"}

    def _extract_xlsx_metadata(self, filepath: Path) -> Dict[str, Any]:
        """Extract Excel metadata"""
        try:
            import openpyxl

            wb = openpyxl.load_workbook(str(filepath), data_only=True)
            props = wb.properties

            return {
                "spreadsheet_info": {
                    "title": props.title,
                    "creator": props.creator,
                    "created": props.created.isoformat() if props.created else None,
                    "modified": props.modified.isoformat() if props.modified else None,
                    "sheet_names": wb.sheetnames,
                    "sheet_count": len(wb.worksheets),
                }
            }
        except ImportError:
            return {"error": "openpyxl not available"}
        except Exception as e:
            return {"error": f"Excel extraction failed: {e}"}

    def _extract_pptx_metadata(self, filepath: Path) -> Dict[str, Any]:
        """Extract PowerPoint metadata"""
        try:
            from pptx import Presentation

            prs = Presentation(str(filepath))
            props = prs.core_properties

            return {
                "presentation_info": {
                    "title": props.title,
                    "author": props.author,
                    "created": props.created.isoformat() if props.created else None,
                    "modified": props.modified.isoformat() if props.modified else None,
                    "slide_count": len(prs.slides),
                }
            }
        except ImportError:
            return {"error": "python-pptx not available"}
        except Exception as e:
            return {"error": f"PowerPoint extraction failed: {e}"}


# Simple usage
def extract_metadata(filepath: Union[str, Path]) -> Dict[str, Any]:
    """Convenience function"""
    extractor = MetadataExtractor()
    return extractor.extract(filepath)


# Example usage
if __name__ == "__main__":
    # Test with actual files
    test_files = [
        r"C:\Users\UserX\Desktop\Github-Workspace\PaperTrail\samples\Adobe Scan Aug 23, 2025 (1).pdf",
        r"C:\Users\UserX\Desktop\Github-Workspace\PaperTrail\samples\Adobe Scan Aug 23, 2025 (3).pdf",
        r"C:\Users\UserX\Desktop\Github-Workspace\PaperTrail\samples\Adobe Scan Aug 23, 2025 (5).pdf",
    ]

    for file_path in test_files:
        if Path(file_path).exists():
            metadata = extract_metadata(file_path)
            print(f"\n--- {file_path} ---")
            print(metadata)
